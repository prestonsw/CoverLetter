# assemble the document from json file

import json 
from docx import Document
import helper

# retrieve json data and convert to dictionary
contents = 'contents.json'

body = json.load(open(contents))


letter = Document('general product.docx') 

bullets = body["bullets"]

print(helper.askForChoices(bullets, body))

bullet_1 = letter.add_paragraph(helper.askForChoices(body["bullets"]))
bullet_2 = letter.add_paragraph(helper.askForChoices(body["bullets"]))
bullet_3 = letter.add_paragraph(helper.askForChoices(body["bullets"]))


letter.save('to-delete.docx')